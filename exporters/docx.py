from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from models.schemas import (
    CertificationEntry,
    ContactFields,
    EducationEntry,
    ExperienceEntry,
    ProjectEntry,
    ResumeBodyJSON,
    SkillGroup,
)
from resume_sections import walk_sections
from templates.library import DEFAULT_TEMPLATE, Template


def _set_font(
    run, name: str = "Calibri", size_pt: float = 11, bold: bool = False
) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _add_section_header(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    _set_font(run, size_pt=11, bold=True)
    # Bottom border to mimic a rule under the heading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_contact_line(doc: Document, contact: ContactFields) -> None:
    """Render the name + contact info at the top of the document."""
    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(contact.name)
    _set_font(name_run, size_pt=16, bold=True)

    # Contact line: email | phone | location | linkedin | github
    parts: list[str] = [contact.email]
    if contact.phone:
        parts.append(contact.phone)
    if contact.location:
        parts.append(contact.location)
    if contact.linkedin:
        parts.append(contact.linkedin)
    if contact.github:
        parts.append(contact.github)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(4)
    contact_run = contact_para.add_run("  |  ".join(parts))
    _set_font(contact_run, size_pt=10)


class _DocxVisitor:
    """Render each resume section into a python-docx Document.

    Ordering, skill-group capping, and optional-section skipping are handled by
    ``walk_sections`` — this visitor only knows how to lay out a section that it
    is asked to render.
    """

    def __init__(self, doc: Document) -> None:
        self._doc = doc

    def summary(self, text: str) -> None:
        _add_section_header(self._doc, "Summary")
        summary_para = self._doc.add_paragraph()
        summary_para.paragraph_format.space_after = Pt(4)
        _set_font(summary_para.add_run(text), size_pt=11)

    def experience(self, entries: list[ExperienceEntry]) -> None:
        _add_section_header(self._doc, "Experience")
        for entry in entries:
            header_para = self._doc.add_paragraph()
            header_para.paragraph_format.space_before = Pt(4)
            header_para.paragraph_format.space_after = Pt(0)
            title_run = header_para.add_run(f"{entry.title}  —  {entry.company}")
            _set_font(title_run, bold=True, size_pt=11)
            if entry.location:
                loc_run = header_para.add_run(f"  |  {entry.location}")
                _set_font(loc_run, size_pt=10)

            dates_para = self._doc.add_paragraph()
            dates_para.paragraph_format.space_after = Pt(1)
            dates_run = dates_para.add_run(f"{entry.start_date} – {entry.end_date}")
            _set_font(dates_run, size_pt=10)
            dates_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

            for bullet in entry.bullets:
                bullet_para = self._doc.add_paragraph(style="List Bullet")
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                bullet_para.paragraph_format.space_after = Pt(1)
                _set_font(bullet_para.add_run(bullet), size_pt=11)

    def skills(self, groups: list[SkillGroup]) -> None:
        skills_line = " | ".join(f"{g.category}: {', '.join(g.skills)}" for g in groups)
        _add_section_header(self._doc, "Skills")
        skills_para = self._doc.add_paragraph()
        skills_para.paragraph_format.space_after = Pt(4)
        _set_font(skills_para.add_run(skills_line), size_pt=11)

    def education(self, entries: list[EducationEntry]) -> None:
        _add_section_header(self._doc, "Education")
        for entry in entries:
            edu_para = self._doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(3)
            edu_para.paragraph_format.space_after = Pt(1)
            degree_run = edu_para.add_run(entry.degree)
            _set_font(degree_run, bold=True, size_pt=11)
            inst_run = edu_para.add_run(
                f"  —  {entry.institution}  |  {entry.graduation_date}"
            )
            _set_font(inst_run, size_pt=10)

    def certifications(self, entries: list[CertificationEntry]) -> None:
        _add_section_header(self._doc, "Certifications")
        for entry in entries:
            cert_para = self._doc.add_paragraph()
            cert_para.paragraph_format.space_before = Pt(3)
            cert_para.paragraph_format.space_after = Pt(1)
            name_run = cert_para.add_run(entry.name)
            _set_font(name_run, bold=True, size_pt=11)
            issuer_text = f"  —  {entry.issuer}"
            if entry.date:
                issuer_text += f"  |  {entry.date}"
            issuer_run = cert_para.add_run(issuer_text)
            _set_font(issuer_run, size_pt=10)

    def projects(self, entries: list[ProjectEntry]) -> None:
        _add_section_header(self._doc, "Projects")
        for entry in entries:
            header_para = self._doc.add_paragraph()
            header_para.paragraph_format.space_before = Pt(4)
            header_para.paragraph_format.space_after = Pt(0)
            name_run = header_para.add_run(entry.name)
            _set_font(name_run, bold=True, size_pt=11)

            desc_para = self._doc.add_paragraph()
            desc_para.paragraph_format.space_after = Pt(1)
            _set_font(desc_para.add_run(entry.description), size_pt=11)

            tech_para = self._doc.add_paragraph()
            tech_para.paragraph_format.space_after = Pt(1)
            tech_run = tech_para.add_run(
                f"Technologies: {', '.join(entry.technologies)}"
            )
            _set_font(tech_run, size_pt=10)
            tech_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

            for bullet in entry.bullets:
                bullet_para = self._doc.add_paragraph(style="List Bullet")
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                bullet_para.paragraph_format.space_after = Pt(1)
                _set_font(bullet_para.add_run(bullet), size_pt=11)


def render(
    resume: ResumeBodyJSON,
    contact: ContactFields,
    template: Template = DEFAULT_TEMPLATE,
) -> bytes:
    """Render a ResumeBodyJSON + ContactFields into DOCX bytes.

    ContactFields are injected here only — they must never appear in LLM prompts.
    The template controls section order and skill group cap.
    """
    doc = Document()

    # Page margins (narrow for more content space)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Default paragraph spacing
    doc.styles["Normal"].paragraph_format.space_after = Pt(2)
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)

    # --- Contact block (PII injected here) ---
    _add_contact_line(doc, contact)

    # --- Sections in template order (traversal + capping shared) ---
    walk_sections(resume, template, _DocxVisitor(doc))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
