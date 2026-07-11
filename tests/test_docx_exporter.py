import io

from docx import Document

from exporters.docx import render
from models.schemas import (
    CertificationEntry,
    ContactFields,
    EducationEntry,
    ExperienceEntry,
    ProjectEntry,
    ResumeBodyJSON,
    SkillGroup,
)
from templates.library import TEMPLATES, Section, Template


def _make_resume(**overrides) -> ResumeBodyJSON:
    defaults = dict(
        summary="A seasoned engineer with cloud experience.",
        rationale="Focused on Python and distributed systems to match the JD.",
        experience=[
            ExperienceEntry(
                title="Senior Engineer",
                company="TechCo",
                location="Remote",
                start_date="Mar 2021",
                end_date="Present",
                bullets=["Led migration to Kubernetes.", "Reduced costs by 20%."],
            )
        ],
        skills=[
            SkillGroup(category="Languages", skills=["Python", "SQL"]),
            SkillGroup(category="Cloud", skills=["Kubernetes", "AWS"]),
        ],
        education=[
            EducationEntry(
                degree="B.S. Computer Science",
                institution="State University",
                graduation_date="May 2018",
            )
        ],
    )
    defaults.update(overrides)
    return ResumeBodyJSON(**defaults)


def _make_contact(**overrides) -> ContactFields:
    defaults = dict(name="Jane Smith", email="jane@example.com")
    defaults.update(overrides)
    return ContactFields(**defaults)


def _full_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Output validity
# ---------------------------------------------------------------------------


def test_render_returns_bytes():
    result = render(_make_resume(), _make_contact())
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_render_produces_valid_docx():
    result = render(_make_resume(), _make_contact())
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


# ---------------------------------------------------------------------------
# Contact fields appear in output
# ---------------------------------------------------------------------------


def test_contact_name_in_docx():
    contact = _make_contact(name="Alice Wonderland", email="alice@example.com")
    doc = Document(io.BytesIO(render(_make_resume(), contact)))
    assert "Alice Wonderland" in _full_text(doc)


def test_contact_email_in_docx():
    contact = _make_contact(email="alice@example.com")
    doc = Document(io.BytesIO(render(_make_resume(), contact)))
    assert "alice@example.com" in _full_text(doc)


def test_optional_contact_fields_in_docx():
    contact = _make_contact(
        phone="+1 555 123 4567",
        location="Austin, TX",
        linkedin="linkedin.com/in/alice",
        github="github.com/alice",
    )
    doc = Document(io.BytesIO(render(_make_resume(), contact)))
    text = _full_text(doc)
    assert "+1 555 123 4567" in text
    assert "Austin, TX" in text
    assert "linkedin.com/in/alice" in text
    assert "github.com/alice" in text


# ---------------------------------------------------------------------------
# Resume content appears in output
# ---------------------------------------------------------------------------


def test_summary_in_docx():
    resume = _make_resume(summary="Unique summary text for testing.")
    doc = Document(io.BytesIO(render(resume, _make_contact())))
    assert "Unique summary text for testing." in _full_text(doc)


def test_experience_title_in_docx():
    doc = Document(io.BytesIO(render(_make_resume(), _make_contact())))
    assert "Senior Engineer" in _full_text(doc)


def test_skills_in_docx():
    doc = Document(io.BytesIO(render(_make_resume(), _make_contact())))
    text = _full_text(doc)
    assert "Python" in text
    assert "Kubernetes" in text


def test_education_in_docx():
    doc = Document(io.BytesIO(render(_make_resume(), _make_contact())))
    assert "State University" in _full_text(doc)


# ---------------------------------------------------------------------------
# Rationale must NOT appear in the rendered document
# ---------------------------------------------------------------------------


def test_rationale_absent_from_docx():
    resume = _make_resume(
        rationale="This rationale should never appear in the DOCX output."
    )
    doc = Document(io.BytesIO(render(resume, _make_contact())))
    assert "This rationale should never appear in the DOCX output." not in _full_text(
        doc
    )


# ---------------------------------------------------------------------------
# Grouped skills rendering
# ---------------------------------------------------------------------------


def test_skills_rendered_with_category_labels():
    doc = Document(io.BytesIO(render(_make_resume(), _make_contact())))
    text = _full_text(doc)
    assert "Languages:" in text
    assert "Cloud:" in text


def test_skills_rendered_as_pipe_separated_groups():
    doc = Document(io.BytesIO(render(_make_resume(), _make_contact())))
    text = _full_text(doc)
    assert "Languages: Python, SQL | Cloud: Kubernetes, AWS" in text


def test_max_skill_groups_cap_respected():
    many_groups = [
        SkillGroup(category=f"Cat{i}", skills=[f"skill{i}"]) for i in range(10)
    ]
    resume = _make_resume(skills=many_groups)
    # standard template has max_skill_groups=5
    doc = Document(io.BytesIO(render(resume, _make_contact(), TEMPLATES["standard"])))
    text = _full_text(doc)
    assert "Cat4:" in text  # 5th group (index 4) present
    assert "Cat5:" not in text  # 6th group absent


# ---------------------------------------------------------------------------
# Certifications renderer
# ---------------------------------------------------------------------------

_CERTS_TEMPLATE = Template(
    id="certs_test",
    name="Certs Test",
    description="Test template including certifications.",
    sections=[Section.SUMMARY, Section.CERTIFICATIONS, Section.EDUCATION],
)


def test_certifications_rendered_when_present():
    resume = _make_resume(
        certifications=[
            CertificationEntry(name="AWS SAA", issuer="Amazon", date="Jun 2023")
        ]
    )
    doc = Document(io.BytesIO(render(resume, _make_contact(), _CERTS_TEMPLATE)))
    text = _full_text(doc)
    assert "AWS SAA" in text
    assert "Amazon" in text
    assert "Jun 2023" in text


def test_certification_no_date_renders_without_pipe():
    resume = _make_resume(
        certifications=[CertificationEntry(name="CKA", issuer="CNCF")]
    )
    doc = Document(io.BytesIO(render(resume, _make_contact(), _CERTS_TEMPLATE)))
    text = _full_text(doc)
    assert "CKA" in text
    assert "CNCF" in text


# ---------------------------------------------------------------------------
# Projects renderer
# ---------------------------------------------------------------------------


def test_projects_rendered_when_present():
    resume = _make_resume(
        projects=[
            ProjectEntry(
                name="DataPipeline",
                description="An ETL pipeline for analytics.",
                technologies=["Python", "Airflow"],
                bullets=["Processed 1M rows daily.", "Reduced runtime by 40%."],
            )
        ]
    )
    technical = TEMPLATES["technical"]
    doc = Document(io.BytesIO(render(resume, _make_contact(), technical)))
    text = _full_text(doc)
    assert "DataPipeline" in text
    assert "Processed 1M rows daily." in text


# ---------------------------------------------------------------------------
# None sections silently skipped
# ---------------------------------------------------------------------------


def test_none_certifications_silently_skipped():
    resume = _make_resume(certifications=None)
    technical = TEMPLATES["technical"]
    doc = Document(io.BytesIO(render(resume, _make_contact(), technical)))
    text = _full_text(doc)
    assert "CERTIFICATIONS" not in text.upper()


def test_none_projects_silently_skipped():
    resume = _make_resume(projects=None)
    technical = TEMPLATES["technical"]
    doc = Document(io.BytesIO(render(resume, _make_contact(), technical)))
    text = _full_text(doc)
    assert "PROJECTS" not in text.upper()


# ---------------------------------------------------------------------------
# Template section order
# ---------------------------------------------------------------------------


def test_template_section_order_respected():
    """Technical template puts Skills before Experience."""
    resume = _make_resume()
    doc = Document(io.BytesIO(render(resume, _make_contact(), TEMPLATES["technical"])))
    # Use exact equality to match only section header paragraphs (not content that
    # happens to contain these words when uppercased).
    texts = [p.text.upper() for p in doc.paragraphs]
    skills_idx = next(i for i, t in enumerate(texts) if t == "SKILLS")
    exp_idx = next(i for i, t in enumerate(texts) if t == "EXPERIENCE")
    assert skills_idx < exp_idx


def test_none_summary_silently_skipped():
    resume = _make_resume(summary=None)
    doc = Document(io.BytesIO(render(resume, _make_contact())))
    text = _full_text(doc)
    assert "SUMMARY" not in text.upper()


def test_none_summary_does_not_raise():
    resume = _make_resume(summary=None)
    result = render(resume, _make_contact())
    assert isinstance(result, bytes)


def test_standard_template_omits_projects():
    resume = _make_resume(
        projects=[
            ProjectEntry(
                name="SideProject",
                description="A cool side project.",
                technologies=["Go"],
                bullets=["Shipped it."],
            )
        ]
    )
    doc = Document(io.BytesIO(render(resume, _make_contact(), TEMPLATES["standard"])))
    text = _full_text(doc)
    assert "SideProject" not in text


def test_technical_template_includes_projects():
    resume = _make_resume(
        projects=[
            ProjectEntry(
                name="SideProject",
                description="A cool side project.",
                technologies=["Go"],
                bullets=["Shipped it."],
            )
        ]
    )
    doc = Document(io.BytesIO(render(resume, _make_contact(), TEMPLATES["technical"])))
    text = _full_text(doc)
    assert "SideProject" in text
